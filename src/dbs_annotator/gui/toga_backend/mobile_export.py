"""Mobile-safe export path for the Toga build.

The Qt build funnels PDF output through
``src/dbs_annotator/utils/session_exporter.py`` which relies on
``docx2pdf`` (Word COM / AppleScript) or LibreOffice -- none of these
work on iOS or Android. This module provides a pure-Python replacement
based on ReportLab plus ``python-docx``, matching the structure of the
desktop reports well enough for a v1 mobile experience.

The functions here accept the same ``dict``-shaped session data the Qt
exporter produces (see ``utils/session_exporter.SessionExporter``) so
both backends can share the upstream collection logic once Phase 2
completes.
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)


def write_session_docx(
    path: Path,
    session: dict[str, Any],
    *,
    electrode_png: Path | None = None,
) -> Path:
    """Write a DOCX report (pure Python, works on every platform)."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    doc = Document()
    title = doc.add_heading("Clinical DBS Session Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = session.get("metadata", {})
    if meta:
        doc.add_heading("Session metadata", level=1)
        for k, v in meta.items():
            doc.add_paragraph(f"{k}: {v}")

    if electrode_png is not None and electrode_png.exists():
        doc.add_heading("Electrode configuration", level=1)
        doc.add_picture(str(electrode_png), width=Inches(2.5))

    for program in session.get("programs", []):
        doc.add_heading(f"Program {program.get('id', '')}", level=1)
        for k, v in program.items():
            if k == "id":
                continue
            doc.add_paragraph(f"{k}: {v}")

    doc.save(str(path))
    return path


def write_session_pdf(
    path: Path,
    session: dict[str, Any],
    *,
    electrode_png: Path | None = None,
) -> Path:
    """Write a PDF report using ReportLab -- no Word / LibreOffice needed."""
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        Image,
        PageBreak,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
    )

    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(str(path), pagesize=LETTER)
    story: list = [
        Paragraph("Clinical DBS Session Report", styles["Title"]),
        Spacer(1, 0.2 * inch),
    ]

    meta = session.get("metadata", {})
    if meta:
        story.append(Paragraph("Session metadata", styles["Heading1"]))
        for k, v in meta.items():
            story.append(Paragraph(f"<b>{k}:</b> {v}", styles["BodyText"]))
        story.append(Spacer(1, 0.2 * inch))

    if electrode_png is not None and electrode_png.exists():
        story.append(Paragraph("Electrode configuration", styles["Heading2"]))
        story.append(Image(str(electrode_png), width=2.5 * inch, height=4.5 * inch))
        story.append(Spacer(1, 0.2 * inch))

    programs = session.get("programs", [])
    for i, program in enumerate(programs):
        if i > 0:
            story.append(PageBreak())
        story.append(Paragraph(f"Program {program.get('id', '')}", styles["Heading1"]))
        for k, v in program.items():
            if k == "id":
                continue
            story.append(Paragraph(f"<b>{k}:</b> {v}", styles["BodyText"]))

    doc.build(story)
    return path
