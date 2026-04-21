"""Mobile-safe report export for the PoC.

Replaces the production app's pipeline in
``src/dbs_annotator/utils/session_exporter.py`` which relies on
``docx2pdf`` (Windows COM / macOS AppleScript) or LibreOffice headless --
none of which run on iOS or Android.

This module writes:

* a ``.docx`` file using ``python-docx`` (pure Python, works everywhere);
* a ``.pdf`` file using ``reportlab`` (pure Python, works everywhere
  including iOS and Android).

Both also embed a PNG of the electrode canvas captured via
``toga.Canvas.as_image()`` from ``canvas_electrode.ElectrodeCanvas.to_png``.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any


def write_docx(
    path: Path, data: dict[str, Any], canvas_png: Path | None = None
) -> Path:
    """Write a minimal DOCX clinical report."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    doc = Document()
    title = doc.add_heading("Clinical DBS Session Report (PoC)", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Electrode model: {data.get('model', '')}")
    doc.add_paragraph(f"Amplitude: {data.get('amplitude_mA', '')} mA")
    doc.add_paragraph(f"Frequency: {data.get('frequency_Hz', '')} Hz")
    doc.add_paragraph(f"Pulse width: {data.get('pulse_width_us', '')} us")
    doc.add_paragraph(f"Clinical scale: {data.get('scale', '')}")

    if canvas_png is not None and canvas_png.exists():
        doc.add_heading("Electrode configuration", level=1)
        doc.add_picture(str(canvas_png), width=Inches(2.5))

    doc.save(str(path))
    return path


def write_pdf(path: Path, data: dict[str, Any], canvas_png: Path | None = None) -> Path:
    """Write a PDF report using ReportLab (no Word/LibreOffice required)."""
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        Image,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
    )

    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(str(path), pagesize=LETTER)
    story: list = [
        Paragraph("Clinical DBS Session Report (PoC)", styles["Title"]),
        Spacer(1, 0.2 * inch),
        Paragraph(f"Electrode model: {data.get('model', '')}", styles["BodyText"]),
        Paragraph(f"Amplitude: {data.get('amplitude_mA', '')} mA", styles["BodyText"]),
        Paragraph(f"Frequency: {data.get('frequency_Hz', '')} Hz", styles["BodyText"]),
        Paragraph(
            f"Pulse width: {data.get('pulse_width_us', '')} us", styles["BodyText"]
        ),
        Paragraph(f"Clinical scale: {data.get('scale', '')}", styles["BodyText"]),
    ]
    if canvas_png is not None and canvas_png.exists():
        story.append(Spacer(1, 0.3 * inch))
        story.append(Paragraph("Electrode configuration", styles["Heading2"]))
        story.append(Image(str(canvas_png), width=2.5 * inch, height=4.5 * inch))

    doc.build(story)
    return path
